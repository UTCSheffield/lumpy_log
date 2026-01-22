"""Custom Markdown to DOCX converter using python-docx.

This module provides a simple but effective markdown to Word document converter
with nice code block handling and no external markdown parsing dependencies.
"""

import re
from pathlib import Path
from typing import List, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
# NEW: imports for mermaid rendering
import os
import tempfile
import urllib.request


def _add_code_block_formatting(paragraph, background_color: Tuple[int, int, int] = (240, 240, 240), apply_text_color: bool = True):
    """Add background color and monospace formatting to a paragraph for code blocks.
    
    Args:
        paragraph: docx paragraph object
        background_color: RGB tuple for background color (default: light gray)
        apply_text_color: If True, apply text color to code blocks (default: True)
    """
    # Set monospace font
    for run in paragraph.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        if apply_text_color:
            run.font.color.rgb = RGBColor(30, 30, 30)
    
    # Add background color using shading
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '%02x%02x%02x' % background_color)
    paragraph._element.get_or_add_pPr().append(shading_elm)
    
    # Set left margin/indent
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.15


# NEW: syntax-highlighted code block rendering (optional pygments)
def _add_highlighted_code_block(doc, code: str, language: str = ""):
    """Add a code block with syntax highlighting using pygments if available."""
    try:
        from pygments import lex
        from pygments.lexers import get_lexer_by_name, TextLexer
        from pygments.styles import get_style_by_name
        from pygments.token import Token
    except Exception:
        # Fallback to plain code block formatting
        para = doc.add_paragraph(code, style='Normal')
        _add_code_block_formatting(para)
        return

    try:
        lexer = get_lexer_by_name(language) if language else TextLexer()
    except Exception:
        lexer = TextLexer()

    style = get_style_by_name('friendly')

    def _parse_style(style_str: str):
        attrs = {'color': None, 'bold': False, 'italic': False}
        if not style_str:
            return attrs
        for part in style_str.split():
            if part == 'bold':
                attrs['bold'] = True
            elif part == 'italic':
                attrs['italic'] = True
            elif part.startswith('#') or re.match(r'^[0-9a-fA-F]{6}$', part):
                hexval = part if part.startswith('#') else f'#{part}'
                attrs['color'] = hexval
        return attrs

    def _style_for(tt):
        t = tt
        while t and t not in style.styles:
            t = t.parent
        return _parse_style(style.styles.get(t, ''))

    # Build paragraphs line-by-line to keep shading per line
    paragraph = doc.add_paragraph(style='Normal')
    _add_code_block_formatting(paragraph, apply_text_color=False)

    for tok_type, tok_val in lex(code, lexer):
        pieces = tok_val.split('\n')
        for idx, piece in enumerate(pieces):
            if piece:
                run = paragraph.add_run(piece)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                attrs = _style_for(tok_type)
                if attrs['bold']:
                    run.bold = True
                if attrs['italic']:
                    run.italic = True
                if attrs['color']:
                    hexval = attrs['color'].lstrip('#')
                    r = int(hexval[0:2], 16)
                    g = int(hexval[2:4], 16)
                    b = int(hexval[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
            if idx < len(pieces) - 1:
                # newline => start a new shaded paragraph
                paragraph = doc.add_paragraph(style='Normal')
                _add_code_block_formatting(paragraph, apply_text_color=False)


# NEW: mermaid rendering via Kroki
def _render_mermaid_and_insert(doc, mermaid_code: str, width_inches: float = 6.0):
    """Render mermaid to PNG via Kroki and insert into the document."""
    try:
        req = urllib.request.Request(
            url='https://kroki.io/mermaid/png',
            data=mermaid_code.encode('utf-8'),
            headers={'Content-Type': 'text/plain'}
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            img_bytes = resp.read()

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        try:
            tmp.write(img_bytes)
            tmp.close()
            doc.add_picture(tmp.name, width=Inches(width_inches))
        finally:
            os.unlink(tmp.name)
    except Exception:
        # Fallback: insert raw code block
        para = doc.add_paragraph(mermaid_code, style='Normal')
        _add_code_block_formatting(para)


def markdown_to_docx(markdown_content: str, output_path: str) -> bool:
    """Convert markdown content to DOCX file.
    
    Args:
        markdown_content: Markdown text to convert
        output_path: Path where to save the DOCX file
    
    Returns:
        True if conversion succeeded, False otherwise
    """
    doc = Document()
    
    # Track if we're in a code block
    in_code_block = False
    code_block_lines = []
    code_language = ""
    
    lines = markdown_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Check for code block start
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                if code_block_lines:
                    code_content = '\n'.join(code_block_lines)
                    lang = (code_language or '').strip().lower()
                    if lang == 'mermaid':
                        _render_mermaid_and_insert(doc, code_content)
                    else:
                        _add_highlighted_code_block(doc, code_content, lang)
                code_block_lines = []
                code_language = ""
            else:
                # Start of code block
                in_code_block = True
                code_language = line.strip()[3:].strip()
                code_block_lines = []
            i += 1
            continue
        
        # If in code block, accumulate lines
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle headers
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            title = header_match.group(2).strip()
            style = f'Heading {level}'
            _add_formatted_paragraph(doc, title, style=style)
            i += 1
            continue
        
        # Handle horizontal rules
        if re.match(r'^[-_*]{3,}$', line.strip()):
            paragraph = doc.add_paragraph()
            pPr = paragraph._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Handle unordered lists
        list_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            list_text = list_match.group(2).strip()
            para = _add_formatted_paragraph(doc, list_text, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            i += 1
            continue
        
        # Handle ordered lists
        ordered_list_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ordered_list_match:
            indent_level = len(ordered_list_match.group(1)) // 2
            list_text = ordered_list_match.group(2).strip()
            para = _add_formatted_paragraph(doc, list_text, style='List Number')
            para.paragraph_format.left_indent = Inches(0.25 + indent_level * 0.25)
            i += 1
            continue
        
        # Handle regular paragraphs
        # Accumulate consecutive non-empty lines as a paragraph
        paragraph_lines = [line.strip()]
        i += 1
        
        while i < len(lines) and lines[i].strip() and not re.match(r'^#+', lines[i]) and not re.match(r'^```', lines[i]) and not re.match(r'^(\s*)[-*+\d]\s', lines[i]):
            paragraph_lines.append(lines[i].strip())
            i += 1
        
        full_text = ' '.join(paragraph_lines)
        if full_text:
            _add_formatted_paragraph(doc, full_text, style='Normal')
    
    # Handle unclosed code block
    if in_code_block and code_block_lines:
        code_content = '\n'.join(code_block_lines)
        lang = (code_language or '').strip().lower()
        if lang == 'mermaid':
            _render_mermaid_and_insert(doc, code_content)
        else:
            _add_highlighted_code_block(doc, code_content, lang)
    
    # Save document
    try:
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error saving DOCX file: {e}")
        return False


def markdown_file_to_docx(markdown_file: str, output_file: str = None) -> bool:
    """Convert a markdown file to DOCX.
    
    Args:
        markdown_file: Path to markdown file
        output_file: Path for output DOCX file (defaults to same name with .docx extension)
    
    Returns:
        True if conversion succeeded, False otherwise
    """
    md_path = Path(markdown_file)
    
    if not md_path.exists():
        print(f"Error: Markdown file not found: {md_path}")
        return False
    
    if output_file is None:
        output_file = str(md_path.with_suffix('.docx'))
    
    try:
        content = md_path.read_text(encoding='utf-8')
        return markdown_to_docx(content, output_file)
    except Exception as e:
        print(f"Error reading markdown file: {e}")
        return False


def _parse_inline_formatting(text: str) -> List[Tuple[str, bool, bool, bool]]:
    """Parse text for bold, italic, and inline code."""
    segments = []
    pos = 0
    for match in re.finditer(r'\*\*(.+?)\*\*|\*(.+?)\*|__(.+?)__|_(.+?)_|`(.+?)`', text):
        if match.start() > pos:
            segments.append((text[pos:match.start()], False, False, False))
        if match.group(1):
            segments.append((match.group(1), True, False, False))
        elif match.group(2):
            segments.append((match.group(2), False, True, False))
        elif match.group(3):
            segments.append((match.group(3), True, False, False))
        elif match.group(4):
            segments.append((match.group(4), False, True, False))
        elif match.group(5):
            segments.append((match.group(5), False, False, True))
        pos = match.end()
    if pos < len(text):
        segments.append((text[pos:], False, False, False))
    return segments if segments else [(text, False, False, False)]


def _add_formatted_paragraph(doc, text: str, style: str = 'Normal', is_code_inline: bool = False):
    """Add a paragraph with inline formatting (bold, italic, code)."""
    paragraph = doc.add_paragraph(style=style)
    segments = _parse_inline_formatting(text)
    for seg_text, is_bold, is_italic, is_inline_code in segments:
        run = paragraph.add_run(seg_text)
        run.bold = is_bold
        run.italic = is_italic
        if is_inline_code or is_code_inline:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(200, 0, 0)
    return paragraph
